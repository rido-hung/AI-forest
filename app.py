# app.py
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

st.set_page_config(page_title="📊 Forest Data Report", layout="wide")

st.title("🌳 Ứng dụng tạo báo cáo điều tra rừng")

uploaded_file = st.file_uploader("📥 Upload file Excel dữ liệu điều tra rừng", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    st.success("✅ Đã tải dữ liệu thành công!")
    
    st.subheader("📋 Dữ liệu gốc")
    st.dataframe(df.head())
    
    # Tính một số thống kê cơ bản
    st.subheader("📊 Thống kê cơ bản")
    stats = df.describe()
    st.dataframe(stats)

    # Vẽ biểu đồ (ví dụ: phân bố đường kính D1.3)
    if "DBH" in df.columns:
        st.subheader("📈 Biểu đồ phân bố đường kính (DBH)")
        fig, ax = plt.subplots()
        df["DBH"].hist(ax=ax, bins=15)
        ax.set_xlabel("Đường kính D1.3 (cm)")
        ax.set_ylabel("Số cây")
        st.pyplot(fig)
    
    # Xuất báo cáo Word
    st.subheader("📝 Xuất báo cáo")
    if st.button("Tạo file Word"):
        doc = Document()
        doc.add_heading("BÁO CÁO ĐIỀU TRA RỪNG", level=1)
        doc.add_paragraph("Kết quả thống kê cơ bản:")
        doc.add_paragraph(stats.to_string())
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📥 Tải về báo cáo Word",
            data=buffer,
            file_name="Forest_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
